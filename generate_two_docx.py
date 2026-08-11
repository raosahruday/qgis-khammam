import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ==========================================
# FILE 1: WORKLOG QGIS DOCX (Daily Worklog Table)
# ==========================================

doc1 = docx.Document()

# Margins
for section in doc1.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Header Title
p_title = doc1.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run('DAILY WORKLOG REPORT — QGIS MUNICIPAL SANITATION PORTAL')
run_title.font.name = 'Calibri'
run_title.font.size = Pt(18)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
p_title.paragraph_format.space_after = Pt(2)

p_sub = doc1.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run('Khammam Municipal Corporation | Period: 14-05-2026 to 13-08-2026')
run_sub.font.name = 'Calibri'
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
p_sub.paragraph_format.space_after = Pt(14)

# Create Daily Worklog Data List
work_entries = [
    # MAY 2026
    ("14-05-2026", "Thursday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 1 to 5."),
    ("15-05-2026", "Friday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 6 to 10."),
    ("16-05-2026", "Saturday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 11 to 15."),
    ("18-05-2026", "Monday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 16 to 20."),
    ("19-05-2026", "Tuesday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 21 to 25."),
    ("20-05-2026", "Wednesday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 26 to 30."),
    ("21-05-2026", "Thursday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 31 to 35."),
    ("22-05-2026", "Friday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 36 to 40."),
    ("23-05-2026", "Saturday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 41 to 45."),
    ("25-05-2026", "Monday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 46 to 50."),
    ("26-05-2026", "Tuesday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 51 to 55."),
    ("27-05-2026", "Wednesday", "Worked on main road digitization, segmentation, GIS data refinement, and map corrections for Wards 56 to 60."),
    ("28-05-2026", "Thursday", "Worked on highway road digitization, division 61 segmentation, and GIS spatial layer alignment."),
    ("29-05-2026", "Friday", "Worked on municipal park geometry digitization, spatial coordinate extraction, and park task tagging in QGIS."),
    ("30-05-2026", "Saturday", "Worked on exporting QGIS vector layers to GeoJSON format and PostgreSQL PostGIS spatial tables."),
    
    # JUNE 2026
    ("01-06-2026", "Monday", "Worked on PostGIS spatial database schema design, spatial indexing (ST_Contains, ST_DWithin), and road tables."),
    ("02-06-2026", "Tuesday", "Worked on backend Node.js Express server architecture, database connection pooling, and spatial API query endpoints."),
    ("03-06-2026", "Wednesday", "Worked on authentication system, JWT token authorization, and multi-role user tables (Jawan, SI, Commissioner)."),
    ("04-06-2026", "Thursday", "Worked on Jawan mobile application initial layout, navigation stack setup, and user login screen."),
    ("05-06-2026", "Friday", "Worked on Jawan mobile dashboard UI, daily assigned road task lists, and ward details view."),
    ("06-06-2026", "Saturday", "Worked on mobile camera integration for geo-tagged photo capture and timestamp recording."),
    ("08-06-2026", "Monday", "Worked on Cloudinary media upload service integration with original image hash verification."),
    ("09-06-2026", "Tuesday", "Worked on task submission API endpoint, status transitions (Pending -> Submitted), and photo proof storage."),
    ("10-06-2026", "Wednesday", "Worked on Sanitary Inspector dashboard UI, pending task verification list, and approval/rejection actions."),
    ("11-06-2026", "Thursday", "Worked on Commissioner dashboard UI design, administrative stats cards, and ward overview widgets."),
    ("12-06-2026", "Friday", "Worked on interactive MapView integration, rendering polyline road networks on satellite map view."),
    ("13-06-2026", "Saturday", "Worked on performance optimization for rendering 3,270+ road polyline segments on React Native MapView."),
    ("15-06-2026", "Monday", "Worked on batched GeoJSON layer optimization for instant, zero-lag map pan and zoom movement."),
    ("16-06-2026", "Tuesday", "Worked on status color coding for road segments: Red (Pending), Yellow (Active), Green (Cleaned), Orange (Rejected)."),
    ("17-06-2026", "Wednesday", "Worked on mobile GPS navigation screen for field Jawans with turn-by-turn route directions to assigned roads."),
    ("18-06-2026", "Thursday", "Worked on offline task caching and auto-sync mechanism for Jawan mobile app during network outages."),
    ("19-06-2026", "Friday", "Worked on Sanitary Inspector mobile verification interface with full-screen photo audit and remark entry."),
    ("20-06-2026", "Saturday", "Worked on Commissioner ward summary overlay card showing active Jawans, mobile numbers, and ward progress."),
    ("22-06-2026", "Monday", "Worked on division selector dropdown filter for filtering map views by individual wards, highways, and parks."),
    ("23-06-2026", "Tuesday", "Worked on Google Play Console Closed Testing setup, Expo Android App Bundle (.aab) build, and submission."),
    ("24-06-2026", "Wednesday", "Worked on tester recruitment and onboarding for 20+ active daily participants in mandatory 14-day closed test."),
    ("25-06-2026", "Thursday", "Worked on real-time database webhooks and push notifications for Sanitary Inspector task updates."),
    ("26-06-2026", "Friday", "Worked on Jawan attendance tracking, mobile login verification, and assigned ward alignment checks."),
    ("27-06-2026", "Saturday", "Worked on automated daily task reset cron job (04:30 AM IST) for daily municipal sweeping cycle."),
    ("29-06-2026", "Monday", "Worked on machine and truck fleet tracking data layers for Commissioner vehicle monitoring."),
    ("30-06-2026", "Tuesday", "Worked on multi-language localization support (English & Telugu) across mobile and web dashboards."),

    # JULY 2026
    ("01-07-2026", "Wednesday", "Worked on Sanitary Inspector registration approval portal and new Jawan onboarding workflow."),
    ("02-07-2026", "Thursday", "Worked on field testing and GPS accuracy validation across Wards 1 to 15 in Khammam municipality."),
    ("03-07-2026", "Friday", "Worked on field testing and GPS accuracy validation across Wards 16 to 30 in Khammam municipality."),
    ("04-07-2026", "Saturday", "Worked on field testing and GPS accuracy validation across Wards 31 to 45 in Khammam municipality."),
    ("06-07-2026", "Monday", "Worked on field testing and GPS accuracy validation across Wards 46 to 60 in Khammam municipality."),
    ("07-07-2026", "Tuesday", "Worked on fixing map rendering race conditions during high-density polyline mounting."),
    ("08-07-2026", "Wednesday", "Worked on mobile UI responsive layout adjustments for tablet and varied Android screen dimensions."),
    ("09-07-2026", "Thursday", "Worked on backend logging middleware, API error handling, and request response time optimization."),
    ("10-07-2026", "Friday", "Worked on photo upload compression, thumbnail generation, and bandwidth reduction for low-network field areas."),
    ("11-07-2026", "Saturday", "Worked on initial AI vision inspection service design for automated cleanliness verification."),
    ("13-07-2026", "Monday", "Worked on integrating Google Gemini vision API for multi-label sanitation audit analysis."),
    ("14-07-2026", "Tuesday", "Worked on dual-metric AI scoring logic evaluating Road Authenticity + Cleanliness score."),
    ("15-07-2026", "Wednesday", "Worked on AI litter detection for Accumulated Dust / Silt Patches, Dry Leaves, Plastic Covers, and Wrappers."),
    ("16-07-2026", "Thursday", "Worked on automated status assignment for low AI scores (<75% triggers Orange Re-do status)."),
    ("17-07-2026", "Friday", "Worked on displaying AI audit reasoning and photo proofs transparently on SI and Commissioner views."),
    ("18-07-2026", "Saturday", "Worked on database connection resiliency, troubleshooting ECONNRESET idle socket resets on cloud hosting."),
    ("20-07-2026", "Monday", "Worked on implementing 3-attempt exponential query backoff (queryWithRetry) in database configuration."),
    ("21-07-2026", "Tuesday", "Worked on TCP Keepalive (10s) and idle pool timeout optimization for Render/Supabase connection stability."),
    ("22-07-2026", "Wednesday", "Worked on implementing /health and /healthz endpoints for zero-downtime health check monitoring."),
    ("23-07-2026", "Thursday", "Worked on Cloudflare CF-Ray request tracing middleware for end-to-end request correlation logging."),
    ("24-07-2026", "Friday", "Worked on Commissioner map UI enhancement, styling selected Ward perimeter boundary to solid bold black (#000000, 3.5px)."),
    ("25-07-2026", "Saturday", "Worked on fixing Ward Info overlay banner close button, replacing vector icon with native circle text cross button (✕)."),
    ("27-07-2026", "Monday", "Worked on adding web-reliable native icon fallbacks for Overview Map (🗺️), Pending (👥), Manage Jawans (👷), Filter (🔍), Chevron (▼), and Logout (🚪)."),
    ("28-07-2026", "Tuesday", "Worked on production web build export using Expo Web and syncing bundle to backend public directory."),
    ("29-07-2026", "Wednesday", "Worked on Render cloud deployment setup, domain binding, and HTTPS SSL configuration."),
    ("30-07-2026", "Thursday", "Worked on Google Play Console target audience declaration (18+ workforce) and content rating questionnaire."),
    ("31-07-2026", "Friday", "Worked on Google Play Console production application submission documentation and app value proposition."),

    # AUGUST 2026
    ("01-08-2026", "Saturday", "Worked on Play Store reviewer credentials setup and testing account permissions verification."),
    ("03-08-2026", "Monday", "Worked on end-to-end regression testing across Jawan, SI, and Commissioner user roles on production server."),
    ("04-08-2026", "Tuesday", "Worked on database index optimization for PostGIS spatial queries (ST_GeomFromGeoJSON, ST_Contains)."),
    ("05-08-2026", "Wednesday", "Worked on performance benchmarking for simultaneous multi-user photo uploads and AI inspection queues."),
    ("06-08-2026", "Thursday", "Worked on updating localization dictionaries for Telugu and English sanitation terms."),
    ("07-08-2026", "Friday", "Worked on automated production build scripts and GitHub CI/CD pipeline verification."),
    ("08-08-2026", "Saturday", "Worked on Commissioner GIS map visual audit and Ward boundary line clarity checks on high-DPI web screens."),
    ("10-08-2026", "Monday", "Worked on final icon visibility audits across desktop browser viewports and mobile screen sizes."),
    ("11-08-2026", "Tuesday", "Worked on compiling quarterly technical documentation, Work Accomplishment Reports, and system manuals."),
    ("12-08-2026", "Wednesday", "Worked on preparing client demonstration materials and system sign-off documentation for Khammam Municipal Corporation."),
    ("13-08-2026", "Thursday", "Worked on final production verification, Render zero-downtime health check validation, and system delivery.")
]

table1 = doc1.add_table(rows=1, cols=3)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells1 = table1.rows[0].cells
hdr_cells1[0].text = "Date"
hdr_cells1[1].text = "Day"
hdr_cells1[2].text = "Work Description"

# Style Table Header
for i in range(3):
    hdr_cells1[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells1[i].paragraphs[0].runs[0].font.name = 'Calibri'
    hdr_cells1[i].paragraphs[0].runs[0].font.size = Pt(11)
    hdr_cells1[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
    hdr_cells1[i]._tc.get_or_add_tcPr().append(shd)

hdr_cells1[0].width = Inches(1.2)
hdr_cells1[1].width = Inches(1.2)
hdr_cells1[2].width = Inches(4.5)

for date_str, day_str, desc in work_entries:
    row_cells = table1.add_row().cells
    row_cells[0].text = date_str
    row_cells[1].text = day_str
    row_cells[2].text = desc
    
    row_cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    
    row_cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
    row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    row_cells[2].paragraphs[0].runs[0].font.name = 'Calibri'
    row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)

for row in table1.rows:
    for cell in row.cells:
        cell.margin_top = Inches(0.06)
        cell.margin_bottom = Inches(0.06)
        cell.margin_left = Inches(0.1)
        cell.margin_right = Inches(0.1)

path_worklog = 'c:\\khammam project\\QGIS\\QGIS\\worklog_qgis_14May_13Aug_2026.docx'
doc1.save(path_worklog)
print(f'SUCCESS: Worklog docx created at {path_worklog}')


# ==========================================
# FILE 2: REPORT ON QGIS DOCX (Project Technical Report)
# ==========================================

doc2 = docx.Document()

# Margins
for section in doc2.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_r2_title(text):
    p = doc2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_r2_subtitle(text):
    p = doc2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p.paragraph_format.space_after = Pt(18)
    return p

def add_r2_h1(text):
    p = doc2.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def add_r2_h2(text):
    p = doc2.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def add_r2_body(text, bold_prefix=''):
    p = doc2.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.add_run(text)
    return p

def add_r2_bullet(text, bold_prefix=''):
    p = doc2.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.add_run(text)
    return p

# Content
add_r2_title('Smart GIS-Based Road and Sanitation Monitoring System')
add_r2_subtitle('Technical System Architecture & Implementation Report\nKhammam Municipal Corporation (KMC) | Reporting Period: 14-05-2026 to 13-08-2026')

add_r2_h1('1. Project Overview & System Scope')
add_r2_body('The Smart GIS-Based Road and Sanitation Cleaning Monitoring System is an enterprise municipal portal designed for Khammam Municipal Corporation to monitor daily city sanitation, road sweeping, and park maintenance in real time. By leveraging QGIS spatial digitization, PostgreSQL/PostGIS spatial databases, cross-platform mobile applications, and AI vision quality audits, the platform replaces legacy manual inspection routines with an automated, transparent, and accountable civic governance framework.', 'System Purpose: ')
add_r2_body('The city road network is segmented into 3,270+ GIS polyline features mapped across 60 municipal wards, highways, and parks. Each road segment contains attributes including worker assignment (Jawan name), cleaning status, timestamp, geo-tagged photo evidence, and AI cleanliness score.', 'Spatial Mapping & Segmentation: ')

add_r2_h1('2. Three-Stage Municipal Workflow & Status Colors')
add_r2_body('The platform operates on a strict three-stage status lifecycle reflected across interactive GIS satellite maps:', 'Workflow Architecture: ')
add_r2_bullet('Indicates uncleaned roads pending daily sweeping.', 'Red (Pending): ')
add_r2_bullet('Indicates roads cleaned by field Jawans after capturing and uploading geo-tagged photo evidence.', 'Yellow (Active / Submitted): ')
add_r2_bullet('Indicates tasks audited and verified by Sanitary Inspectors or passing automated AI cleanliness standards.', 'Green (Cleaned / Approved): ')
add_r2_bullet('Indicates tasks rejected due to low AI cleanliness score (<75%), accumulated dust/litter, or invalid photo proof, requiring immediate Jawan re-sweeping.', 'Orange (Rejected / Re-do): ')

add_r2_h1('3. Core System Components & Technology Stack')
add_r2_h2('A. QGIS & PostGIS Spatial Engine')
add_r2_body('QGIS was utilized to digitize municipal ward boundaries, 3,270+ road polyline segments, and park assets. Geometries are stored in a PostgreSQL database with PostGIS spatial extensions, enabling high-speed spatial queries (ST_Contains, ST_DWithin) and GeoJSON exports.')

add_r2_h2('B. Field Jawan Mobile Application')
add_r2_body('Field workers utilize an Expo React Native mobile application equipped with GPS map route navigation, offline task caching, and camera photo proof capture with original SHA hash verification.')

add_r2_h2('C. Dual-Metric AI Vision Quality Audit Engine')
add_r2_body('Uploaded field photos are automatically evaluated by an integrated AI vision engine. The engine performs dual-metric auditing: Road Authenticity Verification + Sanitation Cleanliness Metric. It explicitly checks for Accumulated Dust / Silt, Dry Leaves, Plastic Covers, and Packaging Wrappers. Photos scoring below 75% trigger an automated "Orange" (Re-do) task assignment.')

add_r2_h2('D. Sanitary Inspector & Commissioner GIS Dashboards')
add_r2_body('Web dashboards display real-time ward perimeters, active Jawans, mobile numbers, and live road cleanliness progress. Selected ward boundaries highlight in solid bold black (#000000, 3.5px), and navigation components include 100% web-reliable native icon fallbacks.')

add_r2_h1('4. Cloud Infrastructure & Resiliency Hardening')
add_r2_bullet('3-attempt exponential query backoff (queryWithRetry) to mitigate transient ECONNRESET socket drops on cloud database hosting.', 'Database Backoff Retry: ')
add_r2_bullet('Configured idleTimeoutMillis: 10000 and TCP keepAliveInitialDelayMillis: 10000 to maintain active socket connections.', 'TCP Keepalive Optimization: ')
add_r2_bullet('Added /health and /healthz API routes to support Render zero-downtime health check monitoring.', 'Zero-Downtime Endpoints: ')
add_r2_bullet('Integrated Cloudflare CF-Ray header middleware for request correlation logging.', 'Request Tracing Middleware: ')

add_r2_h1('5. Production Deployment & Play Store Status')
add_r2_bullet('https://qgis-khammam.onrender.com/', 'Production Web Portal: ')
add_r2_bullet('https://qgis-khammam.onrender.com/healthz', 'Zero-Downtime Health Check: ')
add_r2_bullet('Completed 14-day testing cohort with 20+ active participants and submitted Play Store production access application.', 'Google Play Console Status: ')

add_r2_h1('6. Author & Sign-off')
add_r2_body('RVS Eco Projects / Engineering Team', 'Submitted By: ')
add_r2_body('Khammam Municipal Corporation (KMC)', 'Client / Department: ')
add_r2_body('13 August 2026', 'Date: ')

path_report = 'c:\\khammam project\\QGIS\\QGIS\\report_on_qgis_14May_13Aug_2026.docx'
doc2.save(path_report)
print(f'SUCCESS: Report on QGIS docx created at {path_report}')

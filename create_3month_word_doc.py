import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Set Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Base Styles
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p.paragraph_format.space_after = Pt(18)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def add_body(text, bold_prefix=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.add_run(text)
    return p

def add_bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.add_run(text)
    return p

# Document Content
add_title('QUARTERLY WORK ACCOMPLISHMENT REPORT')
add_subtitle('Khammam Municipal Corporation — Road Sanitation & QGIS Inspection Portal\nReporting Period: 14 May 2026 – 13 August 2026 (3-Month Comprehensive Report)')

# Section 1
add_h1('1. Executive Summary')
add_body('This comprehensive quarterly report synthesizes all software development, spatial data engineering, AI vision integration, database hardening, user experience design, and Google Play Store deployment milestones completed for the Khammam Municipal Sanitation System (Khammam Cleanup) across the 3-month period from May 14, 2026 to August 13, 2026.', 'Project Overview: ')
add_body('Over this period, the project evolved from initial GIS data digitization and mobile app architecture into a production-grade, AI-audited municipal sanitation platform serving Field Jawans, Sanitary Inspectors (SIs), and Municipal Commissioners across 3,270+ road segments and municipal wards.', 'Core Mandate: ')

# Section 2
add_h1('2. Chronological Milestones & Development Phases')

add_h2('Phase 1: Foundation, GIS Digitization & Multi-Role Architecture (14 May – 14 June 2026)')
add_bullet('Digitized and indexed spatial geometries for 3,270+ road segments, ward perimeters, and municipal park assets into a PostGIS PostgreSQL spatial database.', 'QGIS Spatial Mapping: ')
add_bullet('Established secure credential-based role authentication for Jawans, Sanitary Inspectors, Municipal Commissioners, and System Administrators.', 'Multi-Role Portal Setup: ')
add_bullet('Developed the initial Expo React Native cross-platform application supporting mobile field workflows and web administrative monitoring.', 'Cross-Platform Framework: ')

add_h2('Phase 2: Mobile Flow, GeoJSON Optimization & Testing Cohort (15 June – 13 July 2026)')
add_bullet('Refactored frontend map rendering from individual polyline components to high-performance batched GeoJSON layers, eliminating map movement lag.', 'Batched Map Performance: ')
add_bullet('Enforced sequential task reception -> GPS route navigation -> photo proof upload -> submission lock to ensure data integrity.', 'Structured Jawan Workflow: ')
add_bullet('Integrated Cloudinary API with original photo hash verification to ensure authentic, tamper-proof field image uploads.', 'Cloud Media Storage: ')
add_bullet('Initiated Google Play Store Closed Testing track and recruited 20+ active daily testers for mandatory 14-day compliance.', 'Closed Testing Track: ')

add_h2('Phase 3: AI Vision Audit, Resiliency & Production Release (14 July – 13 August 2026)')
add_bullet('Upgraded computer vision service to perform dual-metric analysis: Road Authenticity Verification + Sanitation Cleanliness Metric.', 'Dual-Metric AI Inspection: ')
add_bullet('Audits explicitly for Accumulated Dust / Silt, Dry Leaves, Plastic Covers, and Packaging Wrappers. Scores below 75% trigger automatic "Orange" (Re-do) tasks.', 'Litter Detection & Auto Re-do: ')
add_bullet('Configured 3-attempt exponential query backoff (queryWithRetry) and TCP Keepalive settings (10s) to eliminate ECONNRESET socket drops on Render/Supabase.', 'Database & Cloud Resiliency: ')
add_bullet('Added /health and /healthz routes to enable Render zero-downtime deployment health check monitoring.', 'Zero-Downtime Health Checks: ')
add_bullet('Styled selected ward boundaries as solid bold black (#000000, 3.5px), added native text cross (✕) overlay button, and provided web-reliable icon fallbacks (🗺️, 👥, 👷, 🔍, ▼, 🚪).', 'UI/UX Polish: ')
add_bullet('Completed Play Store target audience and value proposition documentations for Production Access approval.', 'Play Store Production Application: ')

# Section 3
add_h1('3. Core Technical Architecture')

add_bullet('Expo React Native (Web + Android Mobile), GeoJSON Vector Rendering, Custom UI Design Tokens.', 'Frontend Web & Mobile: ')
add_bullet('Node.js, Express, PostgreSQL + PostGIS Extension, Cloudinary Media Engine, Render Cloud Hosting.', 'Backend Infrastructure: ')
add_bullet('Dual-Metric Vision Scoring Engine (Road Type + Multi-Category Litter Analysis).', 'AI Inspection System: ')
add_bullet('GitHub main branch automated build and Render production synchronization.', 'CI/CD & Version Control: ')

# Section 4
add_h1('4. Quarterly Performance & Upgrade Summary')

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_titles = ['System Area', 'Initial Baseline (May 2026)', 'Production State (August 2026)']
for i, title in enumerate(hdr_titles):
    hdr_cells[i].text = title
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = parse_xml(r'<w:shd {} w:fill="0284C7"/>'.format(nsdecls('w')))
    hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

metrics_data = [
    ('GIS Road Coverage', 'Raw shapefiles', '3,270+ PostGIS Digitized & Indexed Road Features'),
    ('Map Polyline Rendering', 'Component lag on zoom/pan', 'Batched GeoJSON (Instant Zero-Lag Rendering)'),
    ('AI Quality Audit', 'Manual inspector verification', 'Automated Dual-Metric AI Audit (<75% Auto Re-do)'),
    ('Database Socket Reliability', 'Intermittent ECONNRESET drops', '100% Stable (TCP Keepalive 10s + 3x Query Retry)'),
    ('Deployment Health', 'Single server restart', 'Zero-Downtime Health Check Endpoints (/healthz)'),
    ('Selected Ward UI', 'Yellow dashed line', 'Solid Bold Black Outline (#000000, 3.5px)'),
    ('Web Icon Rendering', 'Missing vector font glyphs', '100% Reliable Native Icon Fallbacks (🗺️, 👥, 👷, 🔍, ▼, 🚪)'),
    ('Play Store Approval', 'Initial setup', '14-Day Testing Complete & Production Applied')
]

for item in metrics_data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[0].paragraphs[0].runs[0].font.bold = True

for row in table.rows:
    for cell in row.cells:
        cell.margin_top = Inches(0.08)
        cell.margin_bottom = Inches(0.08)
        cell.margin_left = Inches(0.12)
        cell.margin_right = Inches(0.12)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Section 5
add_h1('5. Production Links & Credentials')
add_bullet('https://qgis-khammam.onrender.com/', 'Production Web Dashboard: ')
add_bullet('https://qgis-khammam.onrender.com/healthz', 'Zero-Downtime Health Check: ')
add_bullet('https://github.com/raosahruday/qgis-khammam (Branch: main)', 'GitHub Production Repository: ')

# Section 6
add_h1('6. Author & Sign-off')
add_body('RVS Eco Projects / Engineering Team', 'Submitted By: ')
add_body('Khammam Municipal Corporation (KMC)', 'Client / Department: ')
add_body('13 August 2026', 'Date: ')

output_path = 'c:\\khammam project\\QGIS\\QGIS\\Work_Accomplishment_Report_14May_13Aug_2026.docx'
doc.save(output_path)
print(f'SUCCESS: 3-Month Word document created at {output_path}')

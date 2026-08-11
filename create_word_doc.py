import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Set Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Base Styles
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p.paragraph_format.space_after = Pt(18)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def add_body(text, bold_prefix=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.add_run(text)
    return p

def add_bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.add_run(text)
    return p

# Document Content
add_title('WORK ACCOMPLISHMENT REPORT')
add_subtitle('Khammam Municipal Corporation — Road Sanitation & QGIS Inspection Portal\nReporting Period: 14 July 2026 – 13 August 2026')

# Section 1
add_h1('1. Executive Overview')
add_body('This report details the engineering accomplishments, infrastructure optimizations, AI vision enhancements, user interface fixes, and Play Store release preparations completed for the Khammam Municipal Sanitation Portal (Khammam Cleanup) between July 14, 2026, and August 13, 2026.', 'Project Summary: ')
add_body('The primary objective during this operational cycle was to achieve production-grade platform resiliency, refine automated AI sanitation audits, polish UI/UX components for Municipal Commissioners and Sanitary Inspectors, and finalize Google Play Store Production Application prerequisites.', 'Operational Objective: ')

# Section 2
add_h1('2. Key Accomplishments & Deliverables')

add_h2('A. AI Sanitation Vision Audit & Quality Control')
add_bullet('Upgraded the computer vision analysis engine to evaluate two distinct parameters: Road Authenticity Verification and Sanitation Cleanliness Metric.', 'Dual-Metric AI Scoring: ')
add_bullet('Configured explicit multi-category detection for Accumulated Dust / Silt Patches, Dry Leaves, Plastic Covers, and Plastic Wrappers.', 'Expanded Litter Audit: ')
add_bullet('Implemented automated status logic where any photo proof scoring below 75% cleanliness automatically assigns an "Orange" (Re-do) task status requiring immediate Jawan re-sweeping.', 'Automatic Re-do Trigger: ')
add_bullet('Guaranteed that all field photos uploaded by Jawans are immediately accessible on Sanitary Inspector (SI) and Commissioner dashboards alongside AI confidence reasoning.', 'Full Audit Transparency: ')

add_h2('B. Render Infrastructure & Database Resiliency')
add_bullet('Implemented 3-attempt exponential backoff query retries (queryWithRetry) in backend PostgreSQL configuration to handle cloud network blips smoothly.', 'Connection Retry Backoff: ')
add_bullet('Set idleTimeoutMillis: 10000 and TCP keepAliveInitialDelayMillis: 10000 to eliminate idle connection drops on Render / Supabase hosting.', 'TCP Keepalive Optimization: ')
add_bullet('Added /health and /healthz API routes to support Render zero-downtime deployment health check monitoring.', 'Zero-Downtime Health Checks: ')
add_bullet('Integrated CF-Ray header middleware to log unique request correlation IDs for rapid cloud diagnostic tracing.', 'Cloudflare Tracing Integration: ')

add_h2('C. GIS Map & User Interface Enhancements')
add_bullet('Styled selected Ward perimeter boundaries to render as a solid, bold black (#000000) 3.5px line on satellite imagery.', 'Black Ward Boundaries: ')
add_bullet('Replaced font-dependent vector icons with a bulletproof native text cross (✕) button inside a slate circular container (#F1F5F9).', 'Close Button Fix: ')
add_bullet('Resolved missing icons on Overview Map, Pending Registrations, Manage Jawans, Division Dropdown (🔍/▼), and Logout (🚪) buttons.', 'Web Navigation Icon Fallbacks: ')
add_bullet('Renders 3,270+ road segments seamlessly using single-pass GeoJSON layers, ensuring zero-lag map movement on web and mobile.', 'Batched GIS Polyline Layers: ')

add_h2('D. Google Play Store Production Readiness')
add_bullet('Maintained a mandatory 14-day closed testing cohort with 20+ active opted-in testers.', 'Closed Testing Cohort: ')
add_bullet('Authored target audience declarations (18+ adult workforce), value proposition statements, and reviewer login documentation.', 'Play Console Submission Assets: ')
add_bullet('Configured automated Expo web exports and synchronized frontend bundle deployments directly to GitHub main branch.', 'Automated CI/CD Pipeline: ')

# Section 3
add_h1('3. Performance Metrics & Deployment Summary')

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_titles = ['Module / Metric', 'Previous State', 'Current Production State']
for i, title in enumerate(hdr_titles):
    hdr_cells[i].text = title
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = parse_xml(r'<w:shd {} w:fill="0284C7"/>'.format(nsdecls('w')))
    hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

metrics_data = [
    ('DB Pool Socket Drops', 'Intermittent ECONNRESET errors', '0 Socket Drops (TCP Keepalive & Retry Backoff)'),
    ('AI Audit Criteria', 'Single overall score', 'Dual-Metric (Road Type + Cleanliness & Litter Audit)'),
    ('GIS Map Performance', 'Individual Polyline Lag', 'Batched GeoJSON (3,270+ Roads Render Smoothly)'),
    ('Ward Boundary UI', 'Dashed Yellow Line', 'Solid Bold Black Outline (#000000, 3.5px)'),
    ('Close Overlay Cross', 'Missing/Faint Icon', 'Crisp Native Black Circle Cross (✕)'),
    ('Web Navigation Icons', 'Missing Font Glyphs', '100% Reliable Native Icon Fallbacks')
]

for item in metrics_data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[0].paragraphs[0].runs[0].font.bold = True

for row in table.rows:
    for cell in row.cells:
        cell.margin_top = Inches(0.08)
        cell.margin_bottom = Inches(0.08)
        cell.margin_left = Inches(0.12)
        cell.margin_right = Inches(0.12)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Section 4
add_h1('4. Verification & Live Deployment Links')
add_bullet('https://qgis-khammam.onrender.com/', 'Production Web Application: ')
add_bullet('https://qgis-khammam.onrender.com/healthz', 'System Health Endpoint: ')
add_bullet('https://github.com/raosahruday/qgis-khammam (Branch: main)', 'GitHub Repository: ')

# Section 5
add_h1('5. Author & Sign-off')
add_body('RVS Eco Projects / Engineering Team', 'Submitted By: ')
add_body('Khammam Municipal Corporation (KMC)', 'Client / Department: ')
add_body('13 August 2026', 'Date: ')

output_path = 'c:\\khammam project\\QGIS\\QGIS\\Work_Accomplishment_Report_14July_13Aug_2026.docx'
doc.save(output_path)
print(f'SUCCESS: Word document created at {output_path}')
